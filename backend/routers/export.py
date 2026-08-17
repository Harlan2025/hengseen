"""
导出模块 - 合同文档导出
"""
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.security import HTTPBearer
from fastapi.responses import FileResponse
from datetime import datetime
import uuid

from config import settings
from database import supabase
import uuid
from datetime import datetime
from models.schemas import (
    ExportRequest, ExportHistoryItem, ApiResponse,
    ProjectStatus, ExportFormat
)
from utils.auth_utils import get_current_user

router = APIRouter(prefix="/export", tags=["导出"])


# ========== 导出文档 ==========
@router.post("", response_model=ApiResponse)
async def export_document(req: ExportRequest, user_data: dict = Depends(get_current_user)):
    """
    导出合同文档
    支持Word和Markdown格式
    """
    # 验证项目
    project = await get_project_check(user_id=user_data["sub"], project_id=req.project_id)
    if not project:
        raise HTTPException(status_code=404, detail={"code": 1003, "msg": "项目不存在"})
    
    if project["status"] != ProjectStatus.READY_EXPORT.value:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": 4004, "msg": "项目状态不允许导出"}
        )
    
    # 检查支付状态
    payment_check = await check_payment_status(req.project_id, user_data["sub"])
    if not payment_check.get("paid", False):
        raise HTTPException(
            status_code=status.HTTP_402_PAYMENT_REQUIRED,
            detail={"code": 5002, "msg": "需要支付后才能导出合同"}
        )
    
    # 生成导出文件
    export_id = str(uuid.uuid4())
    file_url = await generate_export_file(req.project_id, req.format, export_id)
    
    # 保存导出记录
    now = datetime.utcnow()
    supabase.table("export_files").insert({
        "export_id": export_id,
        "project_id": req.project_id,
        "user_id": user_data["sub"],
        "format": req.format.value,
        "file_url": file_url,
        "exported_at": now.isoformat(),
    }).execute()
    
    return ApiResponse(data={
        "export_id": export_id,
        "file_url": file_url,
        "format": req.format.value
    })


# ========== 导出历史记录 ==========
@router.get("/history/{project_id}", response_model=ApiResponse)
async def export_history(project_id: str, user_data: dict = Depends(get_current_user)):
    """获取导出历史记录"""
    result = supabase.table("export_files").select("*").eq("project_id", project_id).order("exported_at", desc=True).execute()
    
    items = []
    for row in result.data or []:
        items.append(ExportHistoryItem(
            export_id=row["export_id"],
            project_id=row["project_id"],
            format=row["format"],
            file_url=row["file_url"],
            exported_at=datetime.fromisoformat(row["exported_at"])
        ))
    
    return ApiResponse(data=items)


# ========== 下载导出文件 ==========
@router.get("/{export_id}/download")
async def download_export(export_id: str, user_data: dict = Depends(get_current_user)):
    """下载导出的文件"""
    result = supabase.table("export_files").select("*").eq("export_id", export_id).single().execute()
    
    if not result.data:
        raise HTTPException(status_code=404, detail={"code": 1003, "msg": "导出文件不存在"})
    
    if result.data.get("user_id") != user_data["sub"]:
        raise HTTPException(status_code=403, detail="无权限下载此文件")
    
    file_url = result.data["file_url"]
    
    # TODO: 从Supabase Storage下载文件
    # 这里简化处理
    return FileResponse(
        path=file_url,
        filename=f"contract_{export_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ========== 辅助函数 ==========
async def get_project_check(user_id: str, project_id: str) -> dict:
    """验证项目归属"""
    result = supabase.table("contract_projects").select("*").eq("project_id", project_id).eq("user_id", user_id).single().execute()
    return result.data if result.data else None


async def check_payment_status(project_id: str, user_id: str) -> dict:
    """检查支付状态"""
    result = supabase.table("payment_orders").select("*").eq("project_id", project_id).eq("user_id", user_id).eq("status", "paid").single().execute()
    return {"paid": bool(result.data)}


async def generate_export_file(project_id: str, format: ExportFormat, export_id: str) -> str:
    """生成导出文件"""
    # 获取合同文本
    contract_result = supabase.table("contract_texts").select("*").eq("project_id", project_id).single().execute()
    contract_text = contract_result.data.get("contract_text", "") if contract_result.data else ""
    
    # 获取风险信息
    risk_notes = contract_result.data.get("risk_notes", []) if contract_result.data else []
    
    # 获取项目信息
    project_result = supabase.table("contract_projects").select("*").eq("project_id", project_id).single().execute()
    project_name = project_result.data.get("name", "合同") if project_result.data else "合同"
    
    if format == ExportFormat.WORD:
        # 生成Word文档
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        doc.add_heading(project_name, 0)
        doc.add_paragraph(contract_text)
        
        if risk_notes:
            doc.add_heading("风险提示", level=1)
            for note in risk_notes:
                doc.add_paragraph(f"⚠️ {note}")
        
        doc.add_paragraph("\n本文档仅供参考，不构成法律意见，重大交易请咨询专业律师。")
        
        # 保存文件
        import os
        output_dir = f"{settings.DOMAIN}/exports"
        os.makedirs(output_dir, exist_ok=True)
        file_path = f"{output_dir}/{export_id}.docx"
        doc.save(file_path)
        
        return file_path
    
    else:  # Markdown
        # 生成Markdown文件
        content = f"# {project_name}\n\n{contract_text}\n\n"
        if risk_notes:
            content += "## 风险提示\n\n"
            for note in risk_notes:
                content += f"⚠️ {note}\n\n"
        content += "---\n\n*本文档仅供参考，不构成法律意见，重大交易请咨询专业律师。*"
        
        import os
        output_dir = f"{settings.DOMAIN}/exports"
        os.makedirs(output_dir, exist_ok=True)
        file_path = f"{output_dir}/{export_id}.md"
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        return file_path


def get_auth_header(credentials: dict = Depends(HTTPBearer())) -> dict:
    """获取认证头"""
    from utils.auth_utils import get_current_user as auth_get_current_user
    return auth_get_current_user(credentials)
